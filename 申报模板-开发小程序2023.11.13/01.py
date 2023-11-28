from docx import Document
import openpyxl
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from datetime import datetime
import os

# 设置字体和字号
def set_font_style(element, font_name, font_size):
    element.rPr = parse_xml(
        f'<w:rPr {nsdecls("w")}>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{font_size * 2}"/>'
        f'</w:rPr>'
    )

# 定义文件夹路径和目标字符串
folder_path = "封闭式/01"
folder_path2 = "封闭式/02"
folder_path3 = "封闭式/07"
target_string = "乐惠稳盈2023年第320期"
target_string2 = "2023年6月12日至2023年12月31日"
target_string3 = "1096"
target_string4 ="6.00%"
target_string5 ="50000"
target_string6 ="2023年5月19日"
target_string7 ="LHWY23290"

# 读取xlsx文件
xlsx_filename = "数据来源.xlsx"
wb = openpyxl.load_workbook(xlsx_filename)
sheet = wb.active

# 获取表头的列索引
header_col1 = "产品系列名称"
header_col2 = "期数"
header_col3 = "发行日期"
header_col4 = "理财期限"
header_col5 = "业绩比较基准" #6%变成了0.06
header_col6 = "发行规模"
header_col7 = "报告日期"
header_col8 = "产品编号"
header_col1_index = None
header_col2_index = None
header_col3_index = None
header_col4_index = None
header_col5_index = None
header_col6_index = None
header_col7_index = None
header_col8_index = None

for col in range(1, sheet.max_column + 1):
    cell_value = sheet.cell(row=1, column=col).value
    if cell_value == header_col1:
        header_col1_index = col
    elif cell_value == header_col2:
        header_col2_index = col
    elif cell_value == header_col3:
        header_col3_index = col
    elif cell_value == header_col4:
        header_col4_index = col
    elif cell_value == header_col5:
        header_col5_index = col
    elif cell_value == header_col6:
        header_col6_index = col
    elif cell_value == header_col7:
        header_col7_index = col

# 遍历每一行数据
for row in range(2, sheet.max_row + 1):
    # 创建新的docx文件
    doc = Document()

    # 生成新的字符串
    new_text1 = f"{sheet.cell(row=row, column=header_col1_index).value} {sheet.cell(row=row, column=header_col2_index).value}"


    # 遍历01/docx文件中的段落
    docx_folder_path = os.path.join(os.getcwd(), folder_path)
    for filename in os.listdir(docx_folder_path):
        if filename.endswith(".docx") and filename.startswith("报告主文件-第xxx期"):
            docx_file_path = os.path.join(docx_folder_path, filename)
            doc = Document(docx_file_path)
            new_text1 = (
                f"{sheet.cell(row=row, column=header_col1_index).value} {sheet.cell(row=row, column=header_col2_index).value}"
            )

            # 遍历docx文件中的段落
            for paragraph in doc.paragraphs:
                if target_string in paragraph.text:
                    # 构建替换后的字符串
                    #  乐惠稳盈2023年第320期
                    new_text1 = (
                        f"{sheet.cell(row=row, column=header_col1_index).value} {sheet.cell(row=row, column=header_col2_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string, new_text1)
                if target_string2 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col3_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string2, new_text)
                if target_string3 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col4_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string3, new_text)
                if target_string4 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col5_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string4, new_text)
                if target_string5 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col6_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string5, new_text)
                if target_string6 in paragraph.text:

                    new_text = (
                        f"{sheet.cell(row=row, column=header_col7_index).value}"
                    )
                    #print(new_text)
                    paragraph.text = paragraph.text.replace(target_string6, new_text)

            # 设置字体样式
            for i, paragraph in enumerate(doc.paragraphs):
                for run in paragraph.runs:
                    if i < 3:
                        run.font.size = Pt(19)
                        run.bold = True
                    else:
                        run.font.name = "Arial"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
                        run.font.size = Pt(16)

            # 保存修改后的docx文件为新文件
        new_filename = f"{new_text1}.docx"
        new_file_path = os.path.join(docx_folder_path, new_filename)
        doc.save(new_file_path)

    # 遍历02/docx文件中的段落
    docx_folder_path = os.path.join(os.getcwd(), folder_path2)
    for filename in os.listdir(docx_folder_path):
        if filename.endswith(".docx") and filename.startswith("可行性报告-第xxx期"):
            docx_file_path = os.path.join(docx_folder_path, filename)
            doc = Document(docx_file_path)
            new_text1 = (
                f"{sheet.cell(row=row, column=header_col1_index).value} {sheet.cell(row=row, column=header_col2_index).value}"
            )

            # 遍历docx文件中的段落
            for paragraph in doc.paragraphs:
                if target_string in paragraph.text:
                    # 构建替换后的字符串
                    #  乐惠稳盈2023年第320期
                    new_text1 = (
                        f"{sheet.cell(row=row, column=header_col1_index).value} {sheet.cell(row=row, column=header_col2_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string, new_text1)
                if target_string2 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col3_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string2, new_text)
                if target_string3 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col4_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string3, new_text)
                if target_string4 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col5_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string4, new_text)
                if target_string5 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col6_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string5, new_text)
                if target_string6 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col7_index).value}"
                    )
                    #print(new_text)
                    paragraph.text = paragraph.text.replace(target_string6, new_text)

            # 设置字体样式
            for i, paragraph in enumerate(doc.paragraphs):
                for run in paragraph.runs:
                    if i < 2:
                        run.font.size = Pt(19)
                        run.bold = True
                    else:
                        run.font.name = "Arial"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
                        run.font.size = Pt(16)

            # 保存修改后的docx文件为新文件
        new_filename = f"{new_text1}.docx"
        new_file_path = os.path.join(docx_folder_path, new_filename)
        doc.save(new_file_path)

    # 遍历07/docx文件中的段落
    docx_folder_path = os.path.join(os.getcwd(), folder_path3)
    for filename in os.listdir(docx_folder_path):
        if filename.endswith(".docx") and filename.startswith("乐惠稳盈说明书-第XXX期"):
            docx_file_path = os.path.join(docx_folder_path, filename)
            doc = Document(docx_file_path)
            new_text1 = (
                f"{sheet.cell(row=row, column=header_col1_index).value} {sheet.cell(row=row, column=header_col2_index).value}"
            )

            # 遍历docx文件中的段落
            for paragraph in doc.paragraphs:
                if target_string in paragraph.text:
                    # 构建替换后的字符串
                    #  乐惠稳盈2023年第320期
                    new_text1 = (
                        f"{sheet.cell(row=row, column=header_col1_index).value} {sheet.cell(row=row, column=header_col2_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string, new_text1)
                if target_string2 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col3_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string2, new_text)
                if target_string3 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col4_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string3, new_text)
                if target_string4 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col5_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string4, new_text)
                if target_string5 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col6_index).value}"
                    )

                    paragraph.text = paragraph.text.replace(target_string5, new_text)
                if target_string6 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col7_index).value}"
                    )
                    #print(new_text)
                    paragraph.text = paragraph.text.replace(target_string6, new_text)
                if target_string7 in paragraph.text:
                    new_text = (
                        f"{sheet.cell(row=row, column=header_col8_index).value}"
                    )
                    #print(new_text)
                    paragraph.text = paragraph.text.replace(target_string7, new_text)

            # 设置字体样式
            for i, paragraph in enumerate(doc.paragraphs):
                for run in paragraph.runs:
                    if i < 2:
                        run.font.size = Pt(19)
                        run.bold = True
                    else:
                        run.font.name = "Arial"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
                        run.font.size = Pt(16)

            # 保存修改后的docx文件为新文件
        new_filename = f"{new_text1}.docx"
        new_file_path = os.path.join(docx_folder_path, new_filename)
        doc.save(new_file_path)
